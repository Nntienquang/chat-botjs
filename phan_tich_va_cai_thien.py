"""
Phân tích file Word và cải thiện chatbot
"""
from docx import Document
import os

def phan_tich_docx(file_path):
    """Phân tích chi tiết file .docx"""
    print("="*70)
    print(f"PHÂN TÍCH FILE: {os.path.basename(file_path)}")
    print("="*70)
    
    doc = Document(file_path)
    
    # Phân tích paragraphs
    print("\n[1] PHÂN TÍCH PARAGRAPHS:")
    print("-" * 70)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    print(f"Tổng số paragraphs: {len(paragraphs)}")
    
    # Phân tích độ dài
    lengths = [len(p) for p in paragraphs]
    if lengths:
        print(f"Độ dài trung bình: {sum(lengths)/len(lengths):.0f} ký tự")
        print(f"Độ dài ngắn nhất: {min(lengths)} ký tự")
        print(f"Độ dài dài nhất: {max(lengths)} ký tự")
    
    # Hiển thị một số paragraphs mẫu
    print("\nMột số paragraphs mẫu:")
    for i, para in enumerate(paragraphs[:5], 1):
        preview = para[:100] + "..." if len(para) > 100 else para
        print(f"  [{i}] {preview}")
    
    # Phân tích tables
    print("\n[2] PHÂN TÍCH TABLES:")
    print("-" * 70)
    print(f"Tổng số tables: {len(doc.tables)}")
    
    for i, table in enumerate(doc.tables, 1):
        print(f"\nTable {i}:")
        print(f"  - Số hàng: {len(table.rows)}")
        print(f"  - Số cột: {len(table.columns) if table.rows else 0}")
        if table.rows:
            # Hiển thị hàng đầu tiên
            first_row = [cell.text.strip() for cell in table.rows[0].cells]
            print(f"  - Hàng đầu: {first_row}")
    
    # Phân tích cấu trúc
    print("\n[3] PHÂN TÍCH CẤU TRÚC:")
    print("-" * 70)
    
    # Tìm các tiêu đề (paragraphs ngắn, có thể là tiêu đề)
    titles = [p for p in paragraphs if len(p) < 100 and (p.isupper() or p.endswith(':'))]
    print(f"Số lượng tiêu đề có thể: {len(titles)}")
    if titles:
        print("Một số tiêu đề:")
        for title in titles[:5]:
            print(f"  - {title}")
    
    # Phân tích từ khóa
    print("\n[4] TỪ KHÓA QUAN TRỌNG:")
    print("-" * 70)
    all_text = " ".join(paragraphs).lower()
    
    # Tìm các từ xuất hiện nhiều (có thể là từ khóa)
    words = all_text.split()
    word_freq = {}
    for word in words:
        if len(word) > 3:  # Bỏ qua từ quá ngắn
            word_freq[word] = word_freq.get(word, 0) + 1
    
    top_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:20]
    print("20 từ xuất hiện nhiều nhất:")
    for word, count in top_words:
        print(f"  - {word}: {count} lần")
    
    print("\n" + "="*70)
    print("KẾT LUẬN VÀ ĐỀ XUẤT:")
    print("="*70)
    
    total_chars = sum(lengths)
    print(f"Tổng số ký tự: {total_chars:,}")
    print(f"Số chunks đề xuất (800 ký tự/chunk): {total_chars // 800 + 1}")
    print(f"Độ phức tạp: {'Cao' if total_chars > 10000 else 'Trung bình' if total_chars > 5000 else 'Thấp'}")

if __name__ == "__main__":
    doc_folder = "doc"
    if os.path.exists(doc_folder):
        docx_files = [f for f in os.listdir(doc_folder) if f.endswith('.docx')]
        if docx_files:
            for file in docx_files:
                file_path = os.path.join(doc_folder, file)
                phan_tich_docx(file_path)
        else:
            print(f"Không tìm thấy file .docx trong folder {doc_folder}")
    else:
        print(f"Folder {doc_folder} không tồn tại!")

