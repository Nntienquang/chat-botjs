"""
Chatbot đơn giản học từ tài liệu trong folder doc
"""
import os
import glob
import json
from typing import List, Tuple, Optional, Dict
import re
from docx import Document
import PyPDF2
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

# LLM imports - Sử dụng model generative thực sự
try:
    from transformers import pipeline, AutoTokenizer, AutoModelForCausalLM, AutoModelForSeq2SeqLM
    import torch
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False
    print("Warning: transformers not available, using basic mode")


class DocumentChatbot:
    def __init__(self, doc_folder: str = "doc"):
        """
        Khởi tạo chatbot
        
        Args:
            doc_folder: Đường dẫn đến folder chứa tài liệu
        """
        self.doc_folder = doc_folder
        self.documents = []
        self.document_metadata = []
        self.embeddings = None
        self.model = None
        self.llm_model = None
        self.llm_tokenizer = None
        self.llm_pipeline = None
        self.chunks = []
        self.qa_dataset = []  # Dataset Q&A
        
        print("Đang tải mô hình AI...")
        try:
            self.model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
        except:
            self.model = SentenceTransformer('all-MiniLM-L6-v2')
        print("Đã tải mô hình embedding thành công!")
        
        # Tải Q&A dataset
        self.load_qa_dataset()
        
        # Tải LLM thực sự để tổng hợp câu trả lời như GPT
        if LLM_AVAILABLE:
            print("Đang tải LLM (Large Language Model) để tổng hợp câu trả lời...")
            try:
                # Sử dụng model generative nhẹ nhưng mạnh
                # Có thể dùng GPT-2 hoặc model tiếng Việt tương tự
                model_name = "gpt2"  # GPT-2 base model (nhẹ, nhanh)
                
                try:
                    print(f"Đang tải model: {model_name}...")
                    self.llm_tokenizer = AutoTokenizer.from_pretrained(model_name)
                    self.llm_tokenizer.pad_token = self.llm_tokenizer.eos_token
                    
                    # Sử dụng model nhỏ hơn để tiết kiệm bộ nhớ
                    self.llm_model = AutoModelForCausalLM.from_pretrained(
                        model_name,
                        torch_dtype=torch.float16 if torch.cuda.is_available() else torch.float32,
                        device_map="auto" if torch.cuda.is_available() else None
                    )
                    
                    if not torch.cuda.is_available():
                        self.llm_model = self.llm_model.to('cpu')
                    
                    print("✓ Đã tải LLM thành công! Chatbot sẽ trả lời như GPT.")
                except Exception as e:
                    print(f"Không thể tải LLM model đầy đủ: {e}")
                    print("Sẽ sử dụng pipeline thay thế...")
                    try:
                        self.llm_pipeline = pipeline(
                            "text-generation",
                            model=model_name,
                            tokenizer=model_name,
                            max_length=512,
                            device=0 if torch.cuda.is_available() else -1
                        )
                        self.llm_model = None
                        print("✓ Đã tải LLM pipeline thành công!")
                    except Exception as e2:
                        print(f"Không thể tải LLM: {e2}. Sử dụng chế độ tổng hợp thông minh.")
                        self.llm_model = None
                        self.llm_pipeline = None
            except Exception as e:
                print(f"Lỗi khi khởi tạo LLM: {e}. Sử dụng chế độ cơ bản.")
                self.llm_model = None
                self.llm_pipeline = None
        else:
            self.llm_model = None
            self.llm_pipeline = None
    
    def load_qa_dataset(self):
        """Tải Q&A dataset để học câu hỏi và câu trả lời"""
        qa_file = "qa_dataset.json"
        if os.path.exists(qa_file):
            try:
                with open(qa_file, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    self.qa_dataset = data.get('questions', [])
                    print(f"✓ Đã tải {len(self.qa_dataset)} câu hỏi-đáp từ dataset!")
            except Exception as e:
                print(f"Không thể tải Q&A dataset: {e}")
                self.qa_dataset = []
        else:
            print("Không tìm thấy Q&A dataset, sẽ dùng tìm kiếm semantic.")
            self.qa_dataset = []
    
    def find_qa_match(self, question: str) -> Optional[str]:
        """Tìm câu trả lời từ Q&A dataset - cải thiện để trả lời chính xác hơn"""
        if not self.qa_dataset:
            return None
        
        question_lower = question.lower().strip()
        
        # Xử lý câu hỏi ngắn hoặc không rõ ràng
        if len(question_lower) < 3:
            return None
        
        # Chuẩn hóa câu hỏi - xử lý các biến thể
        question_normalized = question_lower
        # Xử lý các từ viết tắt hoặc sai chính tả
        replacements = {
            'ngày sinh': 'năm sinh',
            'sinh năm': 'năm sinh',
            'mất năm': 'năm mất',
            'hàn mạc tử': 'hàn mặc tử',
            'hàn mặc tư': 'hàn mặc tử'
        }
        for old, new in replacements.items():
            if old in question_normalized:
                question_normalized = question_normalized.replace(old, new)
        
        # Tìm kiếm chính xác trước (so sánh cả câu hỏi gốc và đã chuẩn hóa)
        for qa in self.qa_dataset:
            qa_q = qa['question'].lower().strip()
            if qa_q == question_lower or qa_q == question_normalized:
                return qa['answer']
        
        # Tìm kiếm theo từ khóa quan trọng (ưu tiên cao)
        question_words = set(re.findall(r'\b\w{2,}\b', question_lower))
        
        # Từ khóa quan trọng cần match chính xác
        important_keywords = {
            'năm sinh': ['năm sinh', 'sinh năm', 'ngày sinh'],
            'năm mất': ['năm mất', 'mất năm'],
            'tên thật': ['tên thật', 'tên'],
            'quê quán': ['quê quán', 'quê', 'quảng bình'],
            'thể thơ': ['thể thơ', 'thất ngôn', '7 chữ'],
            'tập thơ': ['tập thơ', 'thơ điên']
        }
        
        # Tìm match với từ khóa quan trọng trước
        for keyword_type, variants in important_keywords.items():
            if any(variant in question_lower for variant in variants):
                for qa in self.qa_dataset:
                    qa_q_lower = qa['question'].lower()
                    qa_keywords = set(qa.get('keywords', []))
                    # Kiểm tra xem Q&A có chứa từ khóa quan trọng không
                    if keyword_type in qa_keywords or any(v in qa_q_lower for v in variants):
                        # Kiểm tra thêm từ khóa liên quan
                        if 'hàn mặc tử' in question_lower and 'hàn mặc tử' in qa_q_lower:
                            return qa['answer']
                        elif keyword_type in ['năm sinh', 'năm mất', 'tên thật', 'quê quán']:
                            return qa['answer']
        
        # Tìm kiếm theo từ khóa thông thường
        best_match = None
        best_score = 0
        
        for qa in self.qa_dataset:
            qa_question_lower = qa['question'].lower()
            qa_keywords = set(qa.get('keywords', []))
            
            # Tính điểm tương đồng
            # 1. So sánh từ khóa (quan trọng hơn)
            keyword_match = len(question_words & qa_keywords)
            # 2. So sánh từ trong câu hỏi
            qa_words = set(re.findall(r'\b\w{2,}\b', qa_question_lower))
            question_similarity = len(question_words & qa_words)
            
            # Tính điểm (ưu tiên keyword match)
            score = keyword_match * 3 + question_similarity
            
            # Bonus nếu có từ khóa quan trọng
            if any(kw in question_lower for kw in ['hàn mặc tử', 'mùa xuân chín']):
                if any(kw in qa_question_lower for kw in ['hàn mặc tử', 'mùa xuân chín']):
                    score += 5
            
            # Chỉ trả về nếu điểm đủ cao và có ít nhất 1 từ khóa match
            if score > best_score and (keyword_match > 0 or question_similarity >= 2):
                best_score = score
                best_match = qa['answer']
        
        # Chỉ trả về nếu điểm đủ cao
        if best_score >= 3:
            return best_match
        
        return None
    
    def read_docx(self, file_path: str) -> str:
        """Đọc file .docx - cải thiện để đọc tables và formatting"""
        try:
            doc = Document(file_path)
            text = []
            
            # Đọc tất cả paragraphs
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text.append(para_text)
            
            # Đọc tất cả tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(" | ".join(row_text))
                if table_text:
                    text.append("\n".join(table_text))
            
            return "\n\n".join(text)
        except Exception as e:
            print(f"Lỗi khi đọc {file_path}: {e}")
            return ""
    
    def read_pdf(self, file_path: str) -> str:
        """Đọc file .pdf"""
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
            return "\n".join(text)
        except Exception as e:
            print(f"Lỗi khi đọc {file_path}: {e}")
            return ""
    
    def read_txt(self, file_path: str) -> str:
        """Đọc file .txt"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            print(f"Lỗi khi đọc {file_path}: {e}")
            return ""
    
    def load_documents(self):
        """Đọc tất cả tài liệu từ folder doc"""
        if not os.path.exists(self.doc_folder):
            print(f"Folder {self.doc_folder} không tồn tại!")
            return
        
        print(f"Đang đọc tài liệu từ folder {self.doc_folder}...")
        
        docx_files = glob.glob(os.path.join(self.doc_folder, "*.docx"))
        pdf_files = glob.glob(os.path.join(self.doc_folder, "*.pdf"))
        txt_files = glob.glob(os.path.join(self.doc_folder, "*.txt"))
        
        all_files = docx_files + pdf_files + txt_files
        
        if not all_files:
            print(f"Không tìm thấy tài liệu nào trong folder {self.doc_folder}!")
            return
        
        for file_path in all_files:
            file_name = os.path.basename(file_path)
            print(f"Đang đọc: {file_name}")
            
            if file_path.endswith('.docx'):
                content = self.read_docx(file_path)
            elif file_path.endswith('.pdf'):
                content = self.read_pdf(file_path)
            elif file_path.endswith('.txt'):
                content = self.read_txt(file_path)
            else:
                continue
            
            if content.strip():
                # Tối ưu chunk size dựa trên độ dài nội dung và cấu trúc
                sentences = re.split(r'[.!?]+', content)
                avg_sentence_length = sum(len(s.strip()) for s in sentences if s.strip()) / max(len([s for s in sentences if s.strip()]), 1)
                
                # Phân tích cấu trúc tài liệu
                has_structure = any(marker in content for marker in ['\n\n', 'I.', '1.', '- ', '* '])
                
                if has_structure:
                    # Tài liệu có cấu trúc rõ ràng, giữ nguyên cấu trúc
                    chunk_size = 1200
                    overlap = 200  # Overlap để giữ ngữ cảnh
                elif avg_sentence_length > 200:
                    chunk_size = 1000
                    overlap = 150
                elif avg_sentence_length > 100:
                    chunk_size = 800
                    overlap = 100
                else:
                    chunk_size = 600
                    overlap = 50
                
                chunks = self.split_text_with_overlap(content, chunk_size=chunk_size, overlap=overlap)
                
                # Lưu metadata chi tiết hơn
                for idx, chunk in enumerate(chunks):
                    self.chunks.append(chunk)
                    self.document_metadata.append({
                        'file': file_name,
                        'path': file_path,
                        'chunk_index': idx,
                        'total_chunks': len(chunks),
                        'chunk_size': len(chunk)
                    })
        
        print(f"Đã đọc {len(self.chunks)} đoạn văn bản từ {len(all_files)} file(s)")
        
        if self.chunks:
            print("Đang tạo embeddings (học kỹ từng đoạn văn)...")
            # Tạo embeddings với batch processing để tối ưu
            self.embeddings = self.model.encode(
                self.chunks, 
                show_progress_bar=True,
                batch_size=32,
                convert_to_numpy=True,
                normalize_embeddings=True  # Chuẩn hóa để tìm kiếm tốt hơn
            )
            print(f"✓ Đã học kỹ {len(self.chunks)} đoạn văn bản!")
            print(f"✓ Tổng số từ đã học: {sum(len(chunk.split()) for chunk in self.chunks):,} từ")
            print("✓ Chatbot đã học xong và sẵn sàng trả lời chính xác như ChatGPT-5!")
    
    def split_text_with_overlap(self, text: str, chunk_size: int = 800, overlap: int = 100) -> List[str]:
        """Chia văn bản với overlap để giữ ngữ cảnh - học kỹ hơn"""
        # Chia theo đoạn văn lớn trước
        paragraphs = re.split(r'\n{2,}', text)
        chunks = []
        current_chunk = ""
        previous_end = ""  # Lưu phần cuối của chunk trước để overlap
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Nếu đoạn văn ngắn, thêm vào chunk hiện tại
            if len(current_chunk) + len(para) + 2 < chunk_size:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
            else:
                # Lưu chunk hiện tại
                if current_chunk:
                    chunks.append(current_chunk.strip())
                    # Lưu phần cuối để overlap
                    previous_end = current_chunk[-overlap:] if len(current_chunk) > overlap else current_chunk
                    # Bắt đầu chunk mới với overlap
                    current_chunk = previous_end + "\n\n" + para if previous_end else para
                else:
                    # Đoạn văn quá dài, chia nhỏ hơn
                    if len(para) > chunk_size:
                        sentences = re.split(r'[.!?]+', para)
                        temp_chunk = ""
                        for sentence in sentences:
                            sentence = sentence.strip()
                            if not sentence or len(sentence) < 10:
                                continue
                            if not sentence[-1] in '.!?':
                                sentence += "."
                            
                            if len(temp_chunk) + len(sentence) + 1 < chunk_size:
                                if temp_chunk:
                                    temp_chunk += " " + sentence
                                else:
                                    temp_chunk = sentence
                            else:
                                if temp_chunk:
                                    chunks.append(temp_chunk.strip())
                                    previous_end = temp_chunk[-overlap:] if len(temp_chunk) > overlap else temp_chunk
                                    temp_chunk = previous_end + " " + sentence if previous_end else sentence
                                else:
                                    temp_chunk = sentence
                        if temp_chunk:
                            current_chunk = temp_chunk
                    else:
                        current_chunk = para
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Lọc và làm sạch chunks
        filtered_chunks = []
        for chunk in chunks:
            chunk = chunk.strip()
            if len(chunk) > 50:  # Chỉ giữ chunks có ý nghĩa
                # Loại bỏ overlap thừa ở đầu
                if len(filtered_chunks) > 0 and chunk.startswith(filtered_chunks[-1][-50:]):
                    chunk = chunk[len(filtered_chunks[-1][-50:]):].strip()
                filtered_chunks.append(chunk)
        
        return filtered_chunks
    
    def split_text(self, text: str, chunk_size: int = 800) -> List[str]:
        """Chia nhỏ văn bản thành các đoạn - tối ưu để giữ ngữ cảnh đầy đủ"""
        # Bước 1: Chia theo đoạn văn lớn (dựa trên \n\n hoặc \n)
        paragraphs = re.split(r'\n{2,}', text)
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # Nếu đoạn văn rất ngắn (< 50 ký tự), có thể là tiêu đề hoặc số thứ tự
            if len(para) < 50:
                # Thêm vào chunk hiện tại nếu có, hoặc tạo chunk mới
                if current_chunk and len(current_chunk) < chunk_size * 0.8:
                    current_chunk += "\n\n" + para
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = para
                continue
            
            # Nếu đoạn văn vừa phải, thêm vào chunk hiện tại
            if len(current_chunk) + len(para) + 2 < chunk_size:
                if current_chunk:
                    current_chunk += "\n\n" + para
                else:
                    current_chunk = para
            else:
                # Lưu chunk hiện tại
                if current_chunk:
                    chunks.append(current_chunk.strip())
                
                # Xử lý đoạn văn dài
                if len(para) > chunk_size:
                    # Chia theo câu một cách thông minh
                    sentences = re.split(r'[.!?]+', para)
                    temp_chunk = ""
                    
                    for sentence in sentences:
                        sentence = sentence.strip()
                        if not sentence or len(sentence) < 10:
                            continue
                        
                        # Thêm dấu chấm nếu cần
                        if not sentence[-1] in '.!?':
                            sentence += "."
                        
                        if len(temp_chunk) + len(sentence) + 1 < chunk_size:
                            if temp_chunk:
                                temp_chunk += " " + sentence
                            else:
                                temp_chunk = sentence
                        else:
                            if temp_chunk:
                                chunks.append(temp_chunk.strip())
                            temp_chunk = sentence
                    
                    if temp_chunk:
                        current_chunk = temp_chunk
                    else:
                        current_chunk = ""
                else:
                    current_chunk = para
        
        # Lưu chunk cuối cùng
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        # Lọc bỏ chunks quá ngắn (có thể là lỗi)
        chunks = [chunk for chunk in chunks if len(chunk) > 50]
        
        return chunks
    
    def search(self, query: str, top_k: int = 5) -> List[Tuple[str, float, dict]]:
        """Tìm kiếm thông tin liên quan đến câu hỏi - tối ưu độ chính xác như ChatGPT"""
        if self.embeddings is None or len(self.chunks) == 0:
            return []
        
        # Chuẩn hóa và mở rộng câu hỏi để tìm kiếm tốt hơn
        query_original = query.strip()
        query = query_original.lower()
        
        # Tạo embedding cho câu hỏi với normalize
        query_embedding = self.model.encode(
            [query_original],  # Dùng câu hỏi gốc để giữ ngữ nghĩa
            show_progress_bar=False,
            normalize_embeddings=True,
            convert_to_numpy=True
        )
        
        # Tính similarity với dot product (nhanh hơn và chính xác hơn khi đã normalize)
        similarities = np.dot(query_embedding, self.embeddings.T)[0]
        
        # Lấy nhiều kết quả hơn để có thể lọc tốt
        num_candidates = min(top_k * 5, len(self.chunks))
        top_indices = np.argsort(similarities)[::-1][:num_candidates]
        
        results = []
        for idx in top_indices:
            score = float(similarities[idx])
            # Ngưỡng động dựa trên phân phối điểm số
            if len(results) == 0:
                threshold = 0.15  # Ngưỡng ban đầu thấp
            elif len(results) < 3:
                threshold = max(0.2, score * 0.7)  # Ngưỡng tăng dần
            else:
                threshold = max(0.25, results[2][1] * 0.8)  # Ngưỡng dựa trên kết quả tốt
            
            if score > threshold:
                results.append((
                    self.chunks[idx],
                    score,
                    self.document_metadata[idx]
                ))
        
        # Sắp xếp lại theo điểm số
        results.sort(key=lambda x: x[1], reverse=True)
        
        # Loại bỏ trùng lặp thông minh
        filtered_results = []
        seen_signatures = set()
        
        for chunk, score, metadata in results:
            # Tạo signature từ đầu và cuối chunk để phát hiện trùng lặp tốt hơn
            chunk_start = chunk[:150].lower().strip()
            chunk_end = chunk[-100:].lower().strip() if len(chunk) > 100 else chunk_start
            signature = (chunk_start, chunk_end)
            
            # Kiểm tra xem có quá giống với kết quả đã có không
            is_duplicate = False
            for sig in seen_signatures:
                # So sánh độ tương đồng của signature
                if chunk_start in sig[0] or sig[0] in chunk_start:
                    if len(chunk_start) > 50:  # Chỉ kiểm tra nếu đủ dài
                        is_duplicate = True
                        break
            
            if not is_duplicate:
                seen_signatures.add(signature)
                filtered_results.append((chunk, score, metadata))
                if len(filtered_results) >= top_k:
                    break
        
        return filtered_results
    
    def _is_greeting(self, text: str) -> bool:
        """Kiểm tra xem có phải câu chào hỏi không"""
        greetings = ['hi', 'hello', 'xin chào', 'chào', 'chào bạn', 'hey', 'hế lô', 'chào bot']
        text_lower = text.lower().strip()
        return text_lower in greetings or any(g in text_lower for g in greetings)
    
    def _detect_question_type(self, question: str) -> str:
        """Phát hiện loại câu hỏi - cải thiện để xử lý nhiều loại câu hỏi hơn"""
        question_lower = question.lower()
        
        # Câu hỏi nhận biết (tác giả, tác phẩm)
        if any(word in question_lower for word in ['tên thật', 'năm sinh', 'năm mất', 'quê quán', 'thể thơ', 'sáng tác', 'tập thơ']):
            return 'factual'
        
        # Câu hỏi về khổ/đoạn cụ thể
        if re.search(r'khổ\s*\d+|đoạn\s*\d+', question_lower):
            return 'specific_section'
        
        # Câu hỏi phân tích nghệ thuật
        if any(word in question_lower for word in ['nghệ thuật', 'biện pháp', 'thủ pháp', 'so sánh với', 'ẩn dụ', 'nhân hóa']):
            return 'artistic'
        
        # Câu hỏi cảm nhận
        if any(word in question_lower for word in ['cảm nhận', 'cảm xúc', 'tâm trạng', 'tâm trí', 'nỗi buồn', 'ý nghĩa']):
            return 'feeling'
        
        # Câu hỏi giải thích
        if any(word in question_lower for word in ['gợi', 'mang ý nghĩa', 'thể hiện', 'có tác dụng', 'vì sao', 'tại sao']):
            return 'explanation'
        
        # Câu hỏi so sánh
        if any(word in question_lower for word in ['so sánh', 'khác nhau', 'giống', 'khác', 'khác biệt']):
            return 'comparison'
        
        # Câu hỏi tổng quan
        if any(word in question_lower for word in ['nội dung', 'nói về', 'chủ đề', 'tổng quan', 'tóm tắt']):
            return 'overview'
        
        # Câu hỏi định nghĩa
        if any(word in question_lower for word in ['là gì', 'là ai', 'định nghĩa', 'khái niệm']):
            return 'definition'
        
        # Câu hỏi liệt kê
        if any(word in question_lower for word in ['những gì', 'các', 'danh sách', 'liệt kê']):
            return 'list'
        
        return 'general'
    
    def _handle_common_questions(self, question: str) -> str:
        """Xử lý các câu hỏi thường gặp - suy nghĩ kỹ trước khi trả lời"""
        question_lower = question.lower().strip()
        
        # Câu hỏi về nội dung chính / tóm tắt
        if any(phrase in question_lower for phrase in ['nội dung chính', 'nội dung là gì', 'nói về gì', 'tóm tắt', 'tổng quan']):
            # Tìm kiếm toàn diện hơn
            results = self.search('nội dung chính bài thơ tác giả tác phẩm', top_k=8)
            if results:
                chunks = [chunk for chunk, _, _ in results]
                # Sử dụng LLM để tổng hợp nếu có
                llm_summary = self._synthesize_with_llm("Hãy tóm tắt nội dung chính của tài liệu một cách đầy đủ và mạch lạc", chunks)
                if llm_summary and len(llm_summary) > 100:
                    return self._clean_answer(llm_summary)
                # Nếu không có LLM, tổng hợp thông minh
                return self._create_comprehensive_summary(chunks)
        
        # Câu hỏi về tác giả (tên thật, năm sinh, quê quán)
        if any(phrase in question_lower for phrase in ['tên thật', 'hàn mặc tử', 'tác giả', 'ai viết']):
            if 'tên thật' in question_lower:
                results = self.search('tên thật nguyễn trọng trí', top_k=2)
            elif 'năm sinh' in question_lower or 'năm mất' in question_lower:
                results = self.search('năm sinh năm mất 1912 1940', top_k=2)
            elif 'quê quán' in question_lower or 'quê' in question_lower:
                results = self.search('quê quán quảng bình', top_k=2)
            else:
                results = self.search('tác giả hàn mặc tử', top_k=3)
            
            if results:
                chunks = [chunk for chunk, _, _ in results]
                # Tìm câu trả lời cụ thể
                for chunk in chunks:
                    if any(word in chunk.lower() for word in ['tên thật', 'nguyễn trọng trí', '1912', '1940', 'quảng bình']):
                        # Trích xuất câu trả lời cụ thể
                        sentences = re.split(r'[.!?]+', chunk)
                        for sentence in sentences:
                            sentence = sentence.strip()
                            if len(sentence) > 20:
                                if 'tên thật' in question_lower and 'nguyễn trọng trí' in sentence.lower():
                                    return sentence + "."
                                elif 'năm sinh' in question_lower and '1912' in sentence:
                                    return sentence + "."
                                elif 'năm mất' in question_lower and '1940' in sentence:
                                    return sentence + "."
                                elif 'quê' in question_lower and 'quảng bình' in sentence.lower():
                                    return sentence + "."
                return "\n\n".join([chunk for chunk in chunks[:2]])
        
        # Câu hỏi về thể thơ, tập thơ
        if any(phrase in question_lower for phrase in ['thể thơ', 'tập thơ', 'sáng tác']):
            if 'thể thơ' in question_lower:
                results = self.search('thể thơ thất ngôn 7 chữ', top_k=2)
            elif 'tập thơ' in question_lower:
                results = self.search('tập thơ thơ điên', top_k=2)
            elif 'sáng tác' in question_lower:
                results = self.search('sáng tác hoàn cảnh 1937', top_k=2)
            else:
                results = self.search(question, top_k=3)
            
            if results:
                chunks = [chunk for chunk, _, _ in results]
                # Tìm câu trả lời cụ thể
                for chunk in chunks:
                    if 'thất ngôn' in chunk.lower() or '7 chữ' in chunk.lower() or 'thơ điên' in chunk.lower():
                        sentences = re.split(r'[.!?]+', chunk)
                        for sentence in sentences:
                            if len(sentence.strip()) > 20:
                                if 'thể thơ' in question_lower and ('thất ngôn' in sentence.lower() or '7 chữ' in sentence):
                                    return sentence + "."
                                elif 'tập thơ' in question_lower and 'thơ điên' in sentence.lower():
                                    return sentence + "."
                return "\n\n".join([chunk for chunk in chunks[:2]])
        
        # Câu hỏi về mục đích
        if any(phrase in question_lower for phrase in ['mục đích', 'mục tiêu', 'ý nghĩa']):
            results = self.search('mục đích mục tiêu', top_k=2)
            if results:
                return "\n\n".join([chunk for chunk, _, _ in results])
        
        return None  # Không phải câu hỏi thường gặp
    
    def _is_simple_question(self, text: str) -> bool:
        """Kiểm tra câu hỏi đơn giản"""
        simple_patterns = ['là gì', 'là ai', 'như thế nào', 'tại sao', 'ở đâu', 'khi nào']
        text_lower = text.lower()
        return any(pattern in text_lower for pattern in simple_patterns)
    
    def _extract_keywords(self, question: str) -> List[str]:
        """Trích xuất từ khóa quan trọng từ câu hỏi"""
        # Loại bỏ các từ không quan trọng
        stop_words = ['là', 'gì', 'như', 'thế', 'nào', 'của', 'và', 'với', 'từ', 'trong', 'về', 'cho', 'đến']
        words = question.lower().split()
        keywords = [w for w in words if w not in stop_words and len(w) > 2]
        return keywords
    
    def answer(self, question: str) -> str:
        """Trả lời câu hỏi dựa trên tài liệu - phiên bản thông minh"""
        if not self.chunks:
            return "Xin lỗi, chưa có tài liệu nào được tải vào hệ thống.\n\nVui lòng đảm bảo có file tài liệu (.docx, .pdf, .txt) trong folder 'doc' và khởi động lại chatbot."
        
        question = question.strip()
        
        # Xử lý câu chào hỏi
        if self._is_greeting(question):
            return "Xin chào! Tôi là chatbot học từ tài liệu. Hãy đặt câu hỏi về nội dung trong tài liệu, tôi sẽ trả lời dựa trên những gì tôi đã học được."
        
        # Xử lý câu hỏi ngắn hoặc không rõ ràng (trước khi tìm kiếm)
        question_clean = question.strip()
        if len(question_clean) < 3:
            return "Xin lỗi, câu hỏi của bạn quá ngắn. Vui lòng đặt câu hỏi cụ thể hơn về nội dung trong tài liệu."
        
        # Xử lý các câu chào hỏi ngắn hoặc không rõ ràng
        short_greetings = ['hi', 'hel', 'hello', 'hey', 'hế lô', 'chào', 'xin chào', 'h', 'he']
        if question_clean.lower() in short_greetings or (len(question_clean) <= 5 and not any(char.isdigit() for char in question_clean) and not any(word in question_clean.lower() for word in ['hỏi', 'gì', 'nào', 'ai', 'đâu'])):
            return "Xin chào! Tôi là chatbot học từ tài liệu. Hãy đặt câu hỏi về nội dung trong tài liệu, tôi sẽ trả lời dựa trên những gì tôi đã học được."
        
        # Ưu tiên 1: Tìm trong Q&A dataset trước (nhanh và chính xác nhất)
        qa_answer = self.find_qa_match(question)
        if qa_answer:
            return qa_answer
        
        # Ưu tiên 2: Xử lý câu hỏi thường gặp
        common_answer = self._handle_common_questions(question)
        if common_answer:
            return self._clean_answer(common_answer)
        
        # Xử lý câu hỏi ngắn và cụ thể (như "Khổ 4", "Đoạn 1")
        question_clean = question.strip()
        
        # Phát hiện câu hỏi về khổ/đoạn cụ thể
        khổ_match = re.search(r'khổ\s*(\d+)', question_clean, re.IGNORECASE)
        đoạn_match = re.search(r'đoạn\s*(\d+)', question_clean, re.IGNORECASE)
        
        if khổ_match:
            khổ_num = khổ_match.group(1)
            # Tìm kiếm cụ thể về khổ đó
            search_query = f"Khổ {khổ_num} khổ thơ {khổ_num}"
            results = self.search(search_query, top_k=3)
            if results:
                # Lọc để chỉ lấy thông tin về khổ cụ thể
                filtered_results = []
                for chunk, score, metadata in results:
                    if f'khổ {khổ_num}' in chunk.lower() or f'khổ{khổ_num}' in chunk.lower():
                        filtered_results.append((chunk, score, metadata))
                if filtered_results:
                    results = filtered_results
        elif đoạn_match:
            đoạn_num = đoạn_match.group(1)
            search_query = f"Đoạn {đoạn_num} đoạn văn {đoạn_num}"
            results = self.search(search_query, top_k=3)
        else:
            # Phát hiện loại câu hỏi để tối ưu tìm kiếm
            question_type = self._detect_question_type(question)
            
            # Điều chỉnh top_k dựa trên loại câu hỏi
            if question_type == 'overview':
                top_k = 6
            elif question_type == 'factual':
                top_k = 3  # Câu hỏi nhận biết cần ít kết quả nhưng chính xác
            elif question_type == 'specific_section':
                top_k = 3
            elif question_type == 'artistic' or question_type == 'feeling':
                top_k = 5  # Câu hỏi phân tích cần nhiều thông tin
            elif question_type == 'explanation':
                top_k = 4
            elif question_type == 'comparison':
                top_k = 6
            elif question_type == 'list':
                top_k = 6
            else:
                top_k = 5
            
            # Tìm kiếm thông tin
            results = self.search(question, top_k=top_k)
        
        if not results:
            # Phân tích câu hỏi để đưa ra gợi ý tốt hơn
            question_words = set(re.findall(r'\b\w{3,}\b', question.lower()))
            
            # Tìm các từ khóa có trong tài liệu
            all_text = " ".join(self.chunks).lower()
            doc_words = set(re.findall(r'\b\w{3,}\b', all_text))
            common_words = question_words & doc_words
            
            if common_words:
                suggestions = f"Tôi không tìm thấy thông tin chính xác về '{question}' trong tài liệu.\n\nTuy nhiên, tài liệu có đề cập đến: {', '.join(list(common_words)[:5])}.\n\nBạn có thể thử:\n- Đặt câu hỏi cụ thể hơn về các chủ đề trên\n- Sử dụng từ khóa: {', '.join(list(common_words)[:3])}"
            else:
                suggestions = f"Xin lỗi, tôi không tìm thấy thông tin về '{question}' trong tài liệu.\n\nGợi ý:\n- Hãy đặt câu hỏi về nội dung trong tài liệu\n- Thử hỏi: 'Nội dung chính là gì?' hoặc 'Tài liệu nói về điều gì?'\n- Sử dụng từ khóa từ tài liệu"
            
            return suggestions
        
        # Kiểm tra độ tương đồng của kết quả tốt nhất
        best_score = results[0][1] if results else 0
        
        # Ngưỡng linh hoạt hơn dựa trên số lượng chunks
        if len(self.chunks) < 10:
            threshold = 0.2  # Tài liệu nhỏ, ngưỡng thấp hơn
        elif len(self.chunks) < 50:
            threshold = 0.22
        else:
            threshold = 0.25  # Tài liệu lớn, ngưỡng cao hơn
        
        if best_score < threshold:
            # Phân tích để đưa ra gợi ý thông minh hơn
            question_lower = question.lower()
            
            # Kiểm tra xem có phải câu hỏi về số, tên riêng không
            if any(char.isdigit() for char in question):
                return f"Xin lỗi, tôi không tìm thấy thông tin chính xác về '{question}' trong tài liệu.\n\nCó thể bạn đang hỏi về:\n- Số thứ tự, trang, đoạn văn cụ thể\n- Thông tin chi tiết không có trong tài liệu\n\nHãy thử hỏi về nội dung, ý nghĩa, hoặc chủ đề chính của tài liệu."
            
            # Kiểm tra từ khóa
            question_words = set(re.findall(r'\b\w{3,}\b', question_lower))
            all_text = " ".join(self.chunks).lower()
            doc_words = set(re.findall(r'\b\w{3,}\b', all_text))
            similar_words = [w for w in question_words if any(dw.startswith(w[:3]) or w.startswith(dw[:3]) for dw in doc_words)]
            
            if similar_words:
                return f"Xin lỗi, tôi không tìm thấy thông tin chính xác về '{question}'.\n\nTuy nhiên, tài liệu có đề cập đến các từ liên quan: {', '.join(similar_words[:5])}.\n\nBạn có thể thử:\n- Hỏi cụ thể hơn về: {similar_words[0] if similar_words else 'chủ đề trong tài liệu'}\n- Hoặc hỏi: 'Nội dung chính là gì?' để xem tổng quan"
            else:
                return f"Xin lỗi, tôi không tìm thấy thông tin về '{question}' trong tài liệu.\n\nGợi ý:\n- Hãy đặt câu hỏi về nội dung trong tài liệu\n- Thử: 'Nội dung chính là gì?', 'Tài liệu nói về điều gì?', 'Hãy tóm tắt'\n- Hoặc hỏi về các chủ đề, khái niệm được đề cập trong tài liệu"
        
        # Sử dụng LLM thực sự để tổng hợp câu trả lời như GPT
        chunks = [chunk for chunk, _, _ in results]
        scores = [score for _, score, _ in results]
        
        # Xử lý đặc biệt cho câu hỏi ngắn và cụ thể (như "Khổ 4")
        if khổ_match or đoạn_match:
            target_num = khổ_match.group(1) if khổ_match else đoạn_match.group(1)
            target_type = "khổ" if khổ_match else "đoạn"
            
            # Tìm tất cả thông tin liên quan đến khổ/đoạn cụ thể
            relevant_chunks = []
            for chunk in chunks:
                chunk_lower = chunk.lower()
                # Kiểm tra xem chunk có chứa thông tin về khổ/đoạn cụ thể không
                patterns = [
                    f"{target_type} {target_num}",
                    f"{target_type}{target_num}",
                    f"* {target_type} {target_num}",
                    f"khổ {target_num}:",
                    f"đoạn {target_num}:"
                ]
                if any(pattern in chunk_lower for pattern in patterns):
                    relevant_chunks.append(chunk)
            
            if relevant_chunks:
                # Tổng hợp thông tin về khổ/đoạn cụ thể
                all_sentences = []
                seen_sentences = set()
                
                for chunk in relevant_chunks:
                    # Tìm phần nói về khổ/đoạn cụ thể
                    sentences = re.split(r'[.!?]+', chunk)
                    for sentence in sentences:
                        sentence = sentence.strip()
                        if len(sentence) < 15:
                            continue
                        
                        # Ưu tiên câu có chứa số khổ/đoạn
                        sentence_lower = sentence.lower()
                        has_target = any(pattern in sentence_lower for pattern in patterns)
                        
                        # Loại bỏ lặp lại
                        sig = sentence[:100].lower()
                        if sig not in seen_sentences:
                            seen_sentences.add(sig)
                            # Ưu tiên câu có chứa thông tin về khổ/đoạn
                            if has_target or len(all_sentences) < 3:
                                all_sentences.append((sentence, has_target))
                
                if all_sentences:
                    # Sắp xếp: câu có target trước
                    all_sentences.sort(key=lambda x: (not x[1], x[0]))
                    answer_sentences = [s[0] for s in all_sentences[:8]]
                    
                    # Tạo câu trả lời mạch lạc
                    answer = ". ".join(answer_sentences)
                    if not answer.endswith(('.', '!', '?')):
                        answer += "."
                    
                    # Làm sạch và loại bỏ lặp lại
                    answer = self._clean_answer(answer)
                    return answer
        
        # Ưu tiên sử dụng LLM để tổng hợp (suy nghĩ và tổng hợp)
        llm_answer = self._synthesize_with_llm(question, chunks)
        if llm_answer and len(llm_answer) > 50:  # Đảm bảo câu trả lời có ý nghĩa và đầy đủ
            # Kiểm tra chất lượng câu trả lời
            if not llm_answer.startswith("Xin lỗi") and len(llm_answer.split()) > 10:
                llm_answer = self._clean_answer(llm_answer)
                return llm_answer
        
        # Nếu không có LLM hoặc LLM trả lời không tốt, dùng tổng hợp thông minh
        answer = self._smart_synthesize(question, chunks, scores)
        
        if not answer:
            # Phân tích câu hỏi để đưa ra gợi ý
            question_lower = question.lower()
            
            # Tìm từ khóa trong câu hỏi
            keywords = [w for w in re.findall(r'\b\w{3,}\b', question_lower) 
                       if w not in ['nào', 'gì', 'sao', 'thế', 'được', 'của', 'trong', 'về', 'cho', 'đến']]
            
            if keywords:
                return f"Xin lỗi, tôi không tìm thấy thông tin đủ liên quan về '{', '.join(keywords[:3])}' trong tài liệu.\n\nBạn có thể:\n- Thử hỏi về chủ đề tổng quan: 'Nội dung chính là gì?'\n- Hoặc đặt câu hỏi cụ thể hơn về nội dung trong tài liệu"
            else:
                return "Xin lỗi, tôi không tìm thấy thông tin đủ liên quan đến câu hỏi của bạn.\n\nHãy thử:\n- Hỏi về nội dung chính: 'Nội dung chính là gì?'\n- Hoặc tóm tắt: 'Hãy tóm tắt tài liệu'\n- Hoặc đặt câu hỏi cụ thể về chủ đề trong tài liệu"
        
        # Làm sạch và format câu trả lời
        answer = answer.strip()
        answer = self._clean_answer(answer)
        
        # Điều chỉnh độ dài dựa trên loại câu hỏi
        question_type = self._detect_question_type(question)
        
        if question_type == 'factual':
            # Câu hỏi nhận biết: ngắn gọn, chính xác (tên thật, năm sinh, quê quán)
            max_length = 200
        elif question_type == 'specific_section':
            # Câu hỏi về khổ/đoạn: vừa phải
            max_length = 500
        elif question_type in ['artistic', 'feeling', 'explanation']:
            # Câu hỏi phân tích: có thể dài hơn
            max_length = 1000
        elif question_type == 'comparison':
            max_length = 1200
        else:
            max_length = 1200
        
        if len(answer) > max_length:
            # Cắt ở câu cuối cùng trước max_length
            cut_point = answer[:max_length].rfind('.')
            if cut_point > max_length * 0.6:
                answer = answer[:cut_point + 1]
            else:
                cut_point = answer[:max_length].rfind('\n')
                if cut_point > max_length * 0.5:
                    answer = answer[:cut_point]
                else:
                    answer = answer[:max_length] + "..."
        
        return answer
    
    def _synthesize_with_llm(self, question: str, context_chunks: List[str]) -> Optional[str]:
        """Sử dụng LLM thực sự để tổng hợp câu trả lời như GPT"""
        if not LLM_AVAILABLE:
            return None
        
        try:
            # Kết hợp context một cách thông minh
            context = "\n\n".join(context_chunks[:4])  # Lấy 4 chunks tốt nhất
            
            # Tạo prompt chuẩn ChatGPT-5 style - học kỹ và trả lời chuẩn
            # Kết hợp context một cách thông minh, ưu tiên chunks có điểm cao
            context_sorted = []
            for i, chunk in enumerate(context_chunks[:5]):  # Lấy 5 chunks tốt nhất
                context_sorted.append(f"[Thông tin {i+1}]\n{chunk}")
            
            context_formatted = "\n\n".join(context_sorted)
            
            # Phân tích loại câu hỏi để tạo prompt phù hợp
            question_lower = question.lower()
            
            # Xác định loại câu hỏi
            if any(word in question_lower for word in ['tên thật', 'năm sinh', 'năm mất', 'quê quán']):
                question_category = "NHẬN BIẾT (câu hỏi về thông tin tác giả, tác phẩm)"
                instruction = "Trả lời CHÍNH XÁC và NGẮN GỌN, chỉ nêu thông tin cụ thể được hỏi."
            elif any(word in question_lower for word in ['gợi', 'mang ý nghĩa', 'thể hiện', 'có tác dụng']):
                question_category = "GIẢI THÍCH (câu hỏi về ý nghĩa, tác dụng)"
                instruction = "Giải thích RÕ RÀNG, có ví dụ cụ thể từ tài liệu, phân tích từng khía cạnh."
            elif any(word in question_lower for word in ['cảm nhận', 'cảm xúc', 'tâm trạng']):
                question_category = "CẢM NHẬN (câu hỏi về cảm xúc, tâm trạng)"
                instruction = "Phân tích SÂU SẮC về cảm xúc, tâm trạng, kết hợp với hình ảnh thơ để làm rõ."
            elif any(word in question_lower for word in ['nghệ thuật', 'biện pháp', 'thủ pháp']):
                question_category = "PHÂN TÍCH NGHỆ THUẬT"
                instruction = "Phân tích CHI TIẾT các biện pháp nghệ thuật, nêu tác dụng cụ thể."
            elif any(word in question_lower for word in ['so sánh', 'khác nhau']):
                question_category = "SO SÁNH"
                instruction = "So sánh RÕ RÀNG, chỉ ra điểm giống và khác, có ví dụ cụ thể."
            else:
                question_category = "TỔNG HỢP"
                instruction = "Trả lời ĐẦY ĐỦ, có cấu trúc rõ ràng, tổng hợp thông tin một cách logic."
            
            prompt = f"""Bạn là ChatGPT-5, một trợ lý AI cực kỳ thông minh và chính xác. Nhiệm vụ của bạn là trả lời câu hỏi dựa trên tài liệu một cách CHÍNH XÁC, ĐẦY ĐỦ và TỰ NHIÊN.

LOẠI CÂU HỎI: {question_category}
HƯỚNG DẪN: {instruction}

THÔNG TIN TÀI LIỆU (đã được sắp xếp theo độ liên quan):
{context_formatted}

CÂU HỎI: {question}

QUY TRÌNH TRẢ LỜI (BẮT BUỘC PHẢI LÀM ĐẦY ĐỦ):
Bước 1: Đọc KỸ tất cả thông tin trong tài liệu
Bước 2: Xác định CHÍNH XÁC câu hỏi đang hỏi về điều gì
Bước 3: Tìm TẤT CẢ thông tin liên quan trong tài liệu
Bước 4: PHÂN TÍCH và TỔNG HỢP thông tin một cách logic
Bước 5: Sắp xếp thông tin theo trình tự hợp lý
Bước 6: Viết câu trả lời TỰ NHIÊN, MẠCH LẠC, DỄ HIỂU

YÊU CẦU BẮT BUỘC:
✓ CHÍNH XÁC 100%: Chỉ dùng thông tin trong tài liệu, KHÔNG bịa đặt
✓ ĐẦY ĐỦ: Trả lời TẤT CẢ khía cạnh của câu hỏi
✓ TỰ NHIÊN: Viết như giáo viên giảng bài, KHÔNG copy-paste nguyên văn
✓ MẠCH LẠC: Có cấu trúc rõ ràng, logic, dễ hiểu
✓ SUY NGHĨ: Tổng hợp và phân tích, KHÔNG chỉ liệt kê
✓ KHÔNG LẶP LẠI: Mỗi ý chỉ nói một lần

TRẢ LỜI (suy nghĩ kỹ, phân tích sâu, trả lời chuẩn):"""
            
            # Sử dụng LLM để generate với cấu hình tối ưu
            if self.llm_pipeline:
                # Sử dụng pipeline với cấu hình tốt hơn
                try:
                    result = self.llm_pipeline(
                        prompt,
                        max_length=min(len(prompt.split()) + 300, 1024),
                        num_return_sequences=1,
                        temperature=0.8,  # Tăng temperature để tự nhiên hơn
                        do_sample=True,
                        top_p=0.95,  # Nucleus sampling
                        top_k=50,
                        repetition_penalty=1.2,  # Tránh lặp lại
                        pad_token_id=self.llm_tokenizer.eos_token_id,
                        eos_token_id=self.llm_tokenizer.eos_token_id
                    )
                    generated_text = result[0]['generated_text']
                    # Lấy phần trả lời (sau prompt)
                    answer = generated_text[len(prompt):].strip()
                except Exception as e:
                    print(f"Lỗi pipeline: {e}")
                    return None
                
            elif self.llm_model and self.llm_tokenizer:
                # Sử dụng model trực tiếp với cấu hình tốt hơn
                try:
                    inputs = self.llm_tokenizer(
                        prompt, 
                        return_tensors="pt", 
                        max_length=512, 
                        truncation=True,
                        padding=True
                    )
                    
                    if torch.cuda.is_available():
                        inputs = {k: v.cuda() for k, v in inputs.items()}
                        self.llm_model = self.llm_model.cuda()
                    
                    with torch.no_grad():
                        outputs = self.llm_model.generate(
                            **inputs,
                            max_new_tokens=250,  # Tăng độ dài
                            temperature=0.8,
                            do_sample=True,
                            top_p=0.95,
                            top_k=50,
                            repetition_penalty=1.2,
                            pad_token_id=self.llm_tokenizer.eos_token_id,
                            eos_token_id=self.llm_tokenizer.eos_token_id,
                            no_repeat_ngram_size=3  # Tránh lặp cụm từ
                        )
                    
                    generated_text = self.llm_tokenizer.decode(outputs[0], skip_special_tokens=True)
                    answer = generated_text[len(prompt):].strip()
                except Exception as e:
                    print(f"Lỗi model: {e}")
                    return None
            else:
                return None
            
            # Làm sạch và cải thiện câu trả lời
            if answer:
                # Loại bỏ phần lặp lại và các marker không cần thiết
                answer = answer.split("CÂU HỎI:")[0].strip()
                answer = answer.split("TRẢ LỜI:")[-1].strip()
                answer = answer.split("THÔNG TIN:")[0].strip()
                
                # Loại bỏ các dòng trống thừa
                answer = re.sub(r'\n{3,}', '\n\n', answer)
                
                # Đảm bảo có độ dài hợp lý và có ý nghĩa
                if len(answer) > 30 and len(answer) < 2000:
                    # Đảm bảo kết thúc bằng dấu câu
                    if answer and answer[-1] not in '.!?':
                        answer += "."
                    return answer
            
            return None
        except Exception as e:
            print(f"Lỗi khi sử dụng LLM: {e}")
            return None
    
    def _smart_synthesize(self, question: str, chunks: List[str], scores: List[float]) -> str:
        """Tổng hợp thông minh - học kỹ và trả lời chuẩn như ChatGPT"""
        """Tổng hợp thông minh câu trả lời từ các chunks - sử dụng logic LLM-like"""
        if not chunks:
            return ""
        
        # Sắp xếp theo điểm số
        sorted_data = sorted(zip(chunks, scores), key=lambda x: x[1], reverse=True)
        
        # Lấy chunks tốt nhất (ưu tiên chunks có điểm cao)
        best_chunks = []
        for chunk, score in sorted_data:
            if score > 0.25:
                best_chunks.append((chunk, score))
            if len(best_chunks) >= 4:  # Lấy tối đa 4 chunks tốt nhất
                break
        
        if not best_chunks:
            return ""
        
        # Nếu chỉ có 1 chunk tốt, trả về trực tiếp
        if len(best_chunks) == 1:
            return best_chunks[0][0]
        
        # Phân tích câu hỏi để quyết định cách tổng hợp
        question_lower = question.lower()
        
        # Nếu là câu hỏi tổng quan, tổng hợp tất cả
        if any(word in question_lower for word in ['tóm tắt', 'tổng quan', 'nội dung chính', 'nói về']):
            return self._synthesize_overview(best_chunks)
        
        # Nếu là câu hỏi cụ thể, tìm phần liên quan nhất
        return self._synthesize_specific(question, best_chunks)
    
    def _synthesize_overview(self, chunks_with_scores: List[Tuple[str, float]]) -> str:
        """Tổng hợp cho câu hỏi tổng quan"""
        all_sentences = []
        seen_content = set()
        
        for chunk, score in chunks_with_scores:
            # Chia thành câu
            sentences = re.split(r'[.!?]+', chunk)
            for sentence in sentences:
                sentence = sentence.strip()
                if len(sentence) < 15:
                    continue
                
                # Loại bỏ trùng lặp
                sig = sentence[:80].lower()
                if sig not in seen_content:
                    seen_content.add(sig)
                    all_sentences.append(sentence)
        
        if all_sentences:
            # Sắp xếp và kết hợp
            answer = ". ".join(all_sentences[:12])
            if not answer.endswith(('.', '!', '?')):
                answer += "."
            return answer
        
        return chunks_with_scores[0][0]
    
    def _synthesize_specific(self, question: str, chunks_with_scores: List[Tuple[str, float]]) -> str:
        """Tổng hợp cho câu hỏi cụ thể - học kỹ và trả lời chuẩn như ChatGPT"""
        # Phân tích câu hỏi kỹ hơn
        question_lower = question.lower()
        question_words = set(re.findall(r'\b\w{3,}\b', question_lower))
        
        # Loại bỏ stop words
        stop_words = {'nào', 'gì', 'sao', 'thế', 'được', 'của', 'trong', 'về', 'cho', 'đến', 'và', 'với', 'từ'}
        question_keywords = question_words - stop_words
        
        best_sentences = []
        seen_content = set()
        
        # Xử lý từng chunk với điểm số
        for chunk, score in chunks_with_scores:
            # Chia thành câu một cách thông minh
            sentences = re.split(r'[.!?]+', chunk)
            
            for sentence in sentences:
                sentence = sentence.strip()
                if len(sentence) < 20:  # Bỏ qua câu quá ngắn
                    continue
                
                # Tính điểm liên quan chi tiết hơn
                sentence_lower = sentence.lower()
                sentence_words = set(re.findall(r'\b\w{3,}\b', sentence_lower))
                
                # Tính điểm dựa trên:
                # 1. Số từ khóa trùng
                keyword_match = len(question_keywords & sentence_words)
                # 2. Độ dài câu (câu dài thường có nhiều thông tin)
                length_score = min(len(sentence) / 200, 1.0)
                # 3. Vị trí trong chunk (câu đầu thường quan trọng hơn)
                
                if keyword_match > 0:
                    relevance = (keyword_match / max(len(question_keywords), 1)) * 0.7 + length_score * 0.3
                    final_score = relevance * score
                    
                    # Chỉ lấy câu có liên quan đáng kể
                    if final_score > 0.15:
                        sig = sentence[:100].lower()
                        if sig not in seen_content:
                            seen_content.add(sig)
                            best_sentences.append((sentence, final_score))
        
        if best_sentences:
            # Sắp xếp theo độ liên quan và điểm số
            best_sentences.sort(key=lambda x: x[1], reverse=True)
            
            # Kết hợp các câu một cách tự nhiên
            # Ưu tiên câu có điểm cao nhất
            top_sentences = [s[0] for s in best_sentences[:10]]
            
            # Tạo câu trả lời mạch lạc
            answer = ". ".join(top_sentences)
            if not answer.endswith(('.', '!', '?')):
                answer += "."
            
            return answer
        
        # Nếu không tìm thấy câu liên quan cụ thể, trả về chunk tốt nhất với format đẹp
        best_chunk = chunks_with_scores[0][0]
        # Làm sạch và format
        best_chunk = re.sub(r'\s+', ' ', best_chunk)
        return best_chunk
    
    def _create_comprehensive_summary(self, chunks: List[str]) -> str:
        """Tạo tóm tắt toàn diện - suy nghĩ và tổng hợp"""
        if not chunks:
            return ""
        
        # Phân tích và tổng hợp thông tin
        summary_parts = []
        
        # Tìm thông tin về tác giả
        author_info = []
        work_info = []
        content_info = []
        
        for chunk in chunks:
            chunk_lower = chunk.lower()
            if any(word in chunk_lower for word in ['hàn mặc tử', 'tác giả', 'nguyễn trọng trí']):
                author_info.append(chunk)
            elif any(word in chunk_lower for word in ['thể thơ', 'xuất xứ', 'hoàn cảnh', 'nhan đề']):
                work_info.append(chunk)
            else:
                content_info.append(chunk)
        
        # Tạo tóm tắt có cấu trúc
        if author_info:
            summary_parts.append("VỀ TÁC GIẢ VÀ TÁC PHẨM:\n" + ". ".join(author_info[:2]) + ".")
        
        if work_info:
            summary_parts.append("\nVỀ TÁC PHẨM:\n" + ". ".join(work_info[:2]) + ".")
        
        if content_info:
            # Tổng hợp nội dung chính
            key_points = []
            for chunk in content_info[:5]:
                # Tìm câu chủ đề
                sentences = re.split(r'[.!?]+', chunk)
                for sentence in sentences:
                    sentence = sentence.strip()
                    if len(sentence) > 30 and any(word in sentence.lower() for word in ['khổ', 'đoạn', 'bức tranh', 'tâm trạng', 'hình ảnh']):
                        key_points.append(sentence)
                        if len(key_points) >= 8:
                            break
                if len(key_points) >= 8:
                    break
            
            if key_points:
                summary_parts.append("\nNỘI DUNG CHÍNH:\n" + ". ".join(key_points) + ".")
        
        if summary_parts:
            return "\n".join(summary_parts)
        
        # Fallback: tổng hợp đơn giản
        return ". ".join(chunks[:4]) + "."
    
    def _clean_answer(self, answer: str) -> str:
        """Làm sạch câu trả lời - loại bỏ lặp lại và format đẹp"""
        if not answer:
            return answer
        
        # Loại bỏ các dòng trống thừa
        answer = re.sub(r'\n{3,}', '\n\n', answer)
        answer = re.sub(r' {2,}', ' ', answer)
        
        # Loại bỏ lặp lại câu một cách thông minh
        sentences = re.split(r'[.!?]+', answer)
        seen_sentences = set()
        unique_sentences = []
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 15:
                continue
            
            # Tạo signature để phát hiện lặp lại (so sánh đầu và cuối)
            sentence_clean = re.sub(r'[^\w\s]', '', sentence.lower())
            sig_start = sentence_clean[:60]
            sig_end = sentence_clean[-40:] if len(sentence_clean) > 40 else sentence_clean
            
            # Kiểm tra xem có quá giống với câu đã có không
            is_duplicate = False
            for seen_sig in seen_sentences:
                # So sánh độ tương đồng
                if sig_start in seen_sig or seen_sig in sig_start:
                    if len(sig_start) > 30:  # Chỉ kiểm tra nếu đủ dài
                        is_duplicate = True
                        break
            
            if not is_duplicate:
                seen_sentences.add(sig_start)
                unique_sentences.append(sentence)
        
        if unique_sentences:
            answer = ". ".join(unique_sentences)
            if not answer.endswith(('.', '!', '?')):
                answer += "."
        else:
            # Nếu tất cả đều trùng, chỉ lấy câu đầu tiên
            if sentences:
                answer = sentences[0].strip() + "."
        
        # Loại bỏ các marker không cần thiết và lặp lại
        answer = re.sub(r'\*+\s*', '', answer)  # Loại bỏ dấu *
        # Loại bỏ "Khổ X:" hoặc "Đoạn X:" lặp lại nhiều lần
        answer = re.sub(r'(Khổ \d+:\s*){2,}', 'Khổ \\1: ', answer, flags=re.IGNORECASE)
        answer = re.sub(r'(Đoạn \d+:\s*){2,}', 'Đoạn \\1: ', answer, flags=re.IGNORECASE)
        
        # Loại bỏ các cụm từ lặp lại
        words = answer.split()
        cleaned_words = []
        prev_word = ""
        for word in words:
            if word != prev_word or len(cleaned_words) == 0:
                cleaned_words.append(word)
            prev_word = word
        answer = " ".join(cleaned_words)
        
        return answer.strip()

