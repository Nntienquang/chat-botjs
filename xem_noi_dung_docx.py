"""
Script để xem nội dung file .docx
"""
from docx import Document
import os

def xem_noi_dung_docx(file_path):
    """Đọc và hiển thị nội dung file .docx"""
    try:
        doc = Document(file_path)
        print("="*60)
        print(f"NỘI DUNG FILE: {os.path.basename(file_path)}")
        print("="*60)
        print()
        
        # Đọc paragraphs
        print("--- PARAGRAPHS ---")
        for i, para in enumerate(doc.paragraphs, 1):
            text = para.text.strip()
            if text:
                print(f"[{i}] {text}")
        
        print()
        print("--- TABLES ---")
        # Đọc tables
        for table_idx, table in enumerate(doc.tables, 1):
            print(f"\nTable {table_idx}:")
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_data.append(cell_text)
                if row_data:
                    print(f"  Row {row_idx + 1}: {' | '.join(row_data)}")
        
        print()
        print("="*60)
        print(f"Tổng số paragraphs: {len([p for p in doc.paragraphs if p.text.strip()])}")
        print(f"Tổng số tables: {len(doc.tables)}")
        print("="*60)
        
    except Exception as e:
        print(f"Lỗi khi đọc file: {e}")

if __name__ == "__main__":
    # Tìm file .docx trong folder doc
    doc_folder = "doc"
    if os.path.exists(doc_folder):
        docx_files = [f for f in os.listdir(doc_folder) if f.endswith('.docx')]
        if docx_files:
            for file in docx_files:
                file_path = os.path.join(doc_folder, file)
                xem_noi_dung_docx(file_path)
        else:
            print(f"Không tìm thấy file .docx trong folder {doc_folder}")
    else:
        print(f"Folder {doc_folder} không tồn tại!")

